from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Adding diagram placeholders with proper formatting...")

doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Function to create a properly formatted diagram box
def add_diagram_placeholder(doc, figure_num, title, width=Inches(5.5), height=Inches(3)):
    """Add a diagram placeholder with figure number, title, and bordered box"""
    
    # Add spacing
    doc.add_paragraph()
    
    # Figure number and title (centered, bold, Times New Roman 11pt)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    
    title_text = f"Figure {figure_num}: {title}"
    title_run = title_para.add_run(title_text)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(11)
    title_run.bold = True
    
    # Add small spacing
    doc.add_paragraph()
    
    # Create table for diagram box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column width
    table.columns[0].width = width
    cell = table.rows[0].cells[0]
    cell.height = height
    
    # Add borders to cell
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'a:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)
    
    # Add placeholder text in center
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_para.paragraph_format.space_before = Pt(0)
    cell_para.paragraph_format.space_after = Pt(0)
    
    placeholder_run = cell_para.add_run("[Insert Diagram Here]")
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add spacing after
    doc.add_paragraph()
    doc.add_paragraph()

# Apply consistent Times New Roman font throughout
print("Applying consistent font formatting...")
for para in doc.paragraphs:
    if para.runs:
        for run in para.runs:
            if not run.font.name or run.font.name in ['Calibri', 'Arial', None]:
                run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

# Now we need to find the right places to insert diagrams
# Let's search for key sections and add diagrams after them

print("Searching for diagram insertion points...")

# We'll add diagrams at the end of relevant sections
# For now, let's add a note section at the end with all diagram requirements

# Find the end of the document (before Bibliography)
bib_index = None
for i, para in enumerate(doc.paragraphs):
    if 'BIBLIOGRAPHY' in para.text.upper():
        bib_index = i
        break

if bib_index:
    # Insert diagram section before Bibliography
    print("Adding diagram requirements section...")
    
    # Add page break
    doc.paragraphs[bib_index].insert_paragraph_before()
    page_break_para = doc.paragraphs[bib_index-1]
    page_break_para.add_run().add_break(WD_BREAK.PAGE)
    
    # Add section title
    diagram_section = doc.paragraphs[bib_index].insert_paragraph_before()
    diagram_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram_title = diagram_section.add_run("DIAGRAMS AND ILLUSTRATIONS")
    diagram_title.font.name = 'Times New Roman'
    diagram_title.font.size = Pt(16)
    diagram_title.bold = True
    
    doc.paragraphs[bib_index].insert_paragraph_before()
    
    # Add note about diagrams
    note_para = doc.paragraphs[bib_index].insert_paragraph_before()
    note_text = """The following diagrams should be inserted at the appropriate locations in the document. Each diagram should be:
    
1. Centered on the page
2. Labeled with "Figure X: [Title]" in Times New Roman, 11pt, Bold, Centered
3. Placed in a bordered box
4. Referenced in the text where appropriate

DIAGRAMS REQUIRED:

Chapter 3.3.1 - Work Breakdown Structure:
- Figure 3.1: Work Breakdown Structure Diagram (Hierarchical tree structure)

Chapter 3.3.2 - Gantt Chart:
- Figure 3.2: Gantt Chart (Project timeline with tasks and milestones)

Chapter 4.1 - UML/DFD:
- Figure 4.1: Entity Relationship Diagram (ERD) - Shows relationships between all entities
- Figure 4.2: Data Flow Diagram - Level 0 (Context Diagram)
- Figure 4.3: Data Flow Diagram - Level 1
- Figure 4.4: Use Case Diagram - Shows actors and use cases
- Figure 4.5: Sequence Diagram - User Registration Flow
- Figure 4.6: Sequence Diagram - Job Application Flow
- Figure 4.7: Activity Diagram - Job Search and Application Process
- Figure 4.8: Component Diagram - System components and relationships
- Figure 4.9: Deployment Diagram - System deployment architecture

Chapter 4.2 - System Flow:
- Figure 4.10: User Registration Flow Diagram
- Figure 4.11: User Login Flow Diagram
- Figure 4.12: Job Browsing and Application Flow Diagram
- Figure 4.13: Job Posting Flow Diagram (Recruiter)
- Figure 4.14: Application Management Flow Diagram
- Figure 4.15: Interview Experience Submission Flow Diagram
- Figure 4.16: Profile Update Flow Diagram
- Figure 4.17: System Architecture Flow Diagram
- Figure 4.18: Authentication Flow Diagram
- Figure 4.19: Error Handling Flow Diagram

Note: Actual diagrams should be created using appropriate tools (Draw.io, Lucidchart, Visio, etc.) and inserted as images at these locations."""
    
    note_para.add_run(note_text)
    for run in note_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Save document
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("\nDocument updated successfully!")
print("\nNext steps:")
print("1. Review the document formatting")
print("2. Insert actual diagrams at the locations mentioned")
print("3. Diagrams should follow the format: Figure X: Title (centered, bold, 11pt)")
print("4. Diagrams should be in bordered boxes, centered on page")

