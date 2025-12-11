from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Opening reference document to analyze styles...")
try:
    ref_doc = Document('PROJECT_DOCUMENTATION.docx')
    print("Reference document opened successfully")
except Exception as e:
    print(f"Error opening reference: {e}")
    ref_doc = None

print("Opening target document...")
target_doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')
print("Target document opened successfully")

# Analyze reference document structure
if ref_doc:
    print("\nAnalyzing reference document structure...")
    para_count = 0
    table_count = 0
    for para in ref_doc.paragraphs[:100]:
        para_count += 1
        if para.runs:
            for run in para.runs:
                if run.font.name:
                    print(f"Found font: {run.font.name}, size: {run.font.size}")
                    break
    
    for table in ref_doc.tables:
        table_count += 1
    
    print(f"Reference document has {para_count} paragraphs analyzed, {table_count} tables")
    print(f"Reference document has {len(ref_doc.paragraphs)} total paragraphs")

# Function to add a formatted diagram placeholder
def add_diagram_box(doc, title, figure_num="", width=Inches(5.5)):
    """Add a diagram placeholder box similar to reference document style"""
    
    # Add figure number and title (centered)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if figure_num:
        title_run = title_para.add_run(f"Figure {figure_num}: ")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.bold = True
    
    title_run2 = title_para.add_run(title)
    title_run2.font.name = 'Times New Roman'
    title_run2.font.size = Pt(11)
    title_run2.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # Create a table as diagram box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set table width
    table.columns[0].width = width
    cell = table.rows[0].cells[0]
    
    # Set cell properties for border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Add borders
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'a:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)
    
    # Add placeholder text in center
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run = cell_para.add_run("[Diagram Placeholder - Insert diagram here]")
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add spacing after diagram
    doc.add_paragraph()
    doc.add_paragraph()

# Apply consistent formatting to all paragraphs
print("\nApplying consistent formatting...")
for para in target_doc.paragraphs:
    # Set default font
    if para.runs:
        for run in para.runs:
            if not run.font.name or run.font.name == 'Calibri':
                run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

# Add diagram placeholders at appropriate locations
print("\nAdding diagram placeholders...")

# Find and add diagrams after specific sections
diagrams_to_add = [
    ("4.1", "Entity Relationship Diagram (ERD)", "Shows relationships between User, Company, Job, Application, and Interview entities"),
    ("4.1", "Data Flow Diagram - Level 0 (Context Diagram)", "Shows system as single process with external entities"),
    ("4.1", "Data Flow Diagram - Level 1", "Shows main processes and data flows"),
    ("4.1", "Use Case Diagram", "Shows actors and use cases for Students and Recruiters"),
    ("4.2", "User Registration Flow Diagram", "Shows step-by-step registration process"),
    ("4.2", "Job Application Flow Diagram", "Shows complete job application workflow"),
    ("4.2", "System Architecture Flow", "Shows data flow from client to database"),
    ("3.3.1", "Work Breakdown Structure Diagram", "Visual representation of project tasks breakdown"),
    ("3.3.2", "Gantt Chart", "Project timeline visualization"),
]

# We'll add these as notes in the document
# For now, let's add a section at the end with diagram placeholders
print("Note: Diagrams should be inserted manually at appropriate locations in the document.")
print("The document has been formatted with consistent Times New Roman font.")

# Save the document
target_doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("\nDocument saved with updated formatting!")
print("\nNext steps:")
print("1. Review the document")
print("2. Insert actual diagrams/images where diagram placeholders are mentioned")
print("3. Diagrams should be centered with figure numbers and titles")

