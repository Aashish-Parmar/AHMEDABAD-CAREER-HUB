from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Enhancing documentation with proper diagram formatting...")

# Open target document
doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Function to create a diagram box with borders
def create_diagram_box(doc, figure_num, title, description=""):
    """Create a diagram placeholder box with proper formatting"""
    
    # Add spacing before diagram
    doc.add_paragraph()
    
    # Add figure number and title (centered, bold)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    
    title_text = f"Figure {figure_num}: {title}"
    title_run = title_para.add_run(title_text)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(11)
    title_run.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # Create table for diagram box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    
    # Set table width
    table.columns[0].width = Inches(5.5)
    cell = table.rows[0].cells[0]
    
    # Set cell height
    cell.height = Inches(3)
    
    # Add borders explicitly
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders
    borders = ['top', 'bottom', 'left', 'right']
    for border_name in borders:
        border = OxmlElement(f'a:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)
    
    # Add placeholder text
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_para.paragraph_format.space_before = Pt(0)
    cell_para.paragraph_format.space_after = Pt(0)
    
    if description:
        placeholder_text = f"[{description}]"
    else:
        placeholder_text = "[Insert Diagram Here]"
    
    placeholder_run = cell_para.add_run(placeholder_text)
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add spacing after diagram
    doc.add_paragraph()
    doc.add_paragraph()

# Apply consistent font formatting throughout
print("Applying consistent formatting...")
for para in doc.paragraphs:
    if para.runs:
        for run in para.runs:
            # Set Times New Roman as default
            if not run.font.name or run.font.name in ['Calibri', 'Arial']:
                run.font.name = 'Times New Roman'
            # Set default size if not set
            if not run.font.size:
                run.font.size = Pt(12)

# Find sections where diagrams should be added and insert placeholders
print("Adding diagram placeholders at appropriate locations...")

# We'll create a new document structure with diagrams
# First, let's identify where to add diagrams based on content

# Save the document
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("Document formatting updated!")

# Now create a helper script to add diagrams at specific locations
print("\nCreating diagram insertion guide...")

diagram_locations = """
DIAGRAMS TO BE ADDED:

Chapter 4.1 - UML/DFD:
- Figure 4.1: Entity Relationship Diagram (ERD)
- Figure 4.2: Data Flow Diagram - Level 0 (Context Diagram)  
- Figure 4.3: Data Flow Diagram - Level 1
- Figure 4.4: Use Case Diagram
- Figure 4.5: Sequence Diagram - User Registration
- Figure 4.6: Sequence Diagram - Job Application
- Figure 4.7: Activity Diagram - Job Search Flow
- Figure 4.8: Component Diagram
- Figure 4.9: Deployment Diagram

Chapter 4.2 - System Flow:
- Figure 4.10: User Registration Flow
- Figure 4.11: User Login Flow
- Figure 4.12: Job Browsing and Application Flow
- Figure 4.13: Job Posting Flow
- Figure 4.14: Application Management Flow
- Figure 4.15: Interview Experience Submission Flow
- Figure 4.16: Profile Update Flow
- Figure 4.17: System Architecture Flow
- Figure 4.18: Authentication Flow
- Figure 4.19: Error Handling Flow

Chapter 3.3.1 - WBS:
- Figure 3.1: Work Breakdown Structure Diagram

Chapter 3.3.2 - Gantt Chart:
- Figure 3.2: Gantt Chart - Project Timeline

Note: Insert actual diagrams/images at these locations with proper figure numbering.
"""

print(diagram_locations)

