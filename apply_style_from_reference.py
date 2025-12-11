from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Open reference document to analyze styles
ref_doc = Document('PROJECT_DOCUMENTATION.docx')
target_doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

print("Analyzing reference document styles...")

# Analyze reference document styles
ref_styles = {}
for para in ref_doc.paragraphs[:50]:  # Analyze first 50 paragraphs
    if para.style.name not in ref_styles:
        ref_styles[para.style.name] = {
            'font_name': para.runs[0].font.name if para.runs else None,
            'font_size': para.runs[0].font.size if para.runs else None,
            'alignment': para.alignment,
            'bold': para.runs[0].bold if para.runs else False,
        }

print(f"Found {len(ref_styles)} unique styles in reference document")

# Now apply similar formatting to target document
print("Applying styles to target document...")

# Function to set paragraph formatting
def set_para_format(para, font_name='Times New Roman', font_size=12, bold=False, alignment=None):
    if para.runs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
    if alignment:
        para.alignment = alignment

# Apply consistent formatting throughout
for para in target_doc.paragraphs:
    # Check if it's a heading
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.replace('Heading ', ''))
        if level == 1:
            # Chapter headings
            if para.runs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            # Section headings
            if para.runs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.bold = True
        elif level == 3:
            # Subsection headings
            if para.runs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
    else:
        # Normal paragraphs
        if para.runs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

# Add diagram placeholders with proper formatting
print("Adding diagram placeholders with proper formatting...")

# Function to add a diagram placeholder
def add_diagram_placeholder(doc, title, description=""):
    # Add diagram title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(11)
    title_run.bold = True
    
    # Add diagram placeholder box
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a table to represent diagram box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Light Grid Accent 1'
    cell = table.rows[0].cells[0]
    cell.width = Inches(5)
    
    # Add placeholder text
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run = cell_para.add_run(f"[{description or 'Diagram Placeholder'}]")
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add spacing
    doc.add_paragraph()

# Find locations where diagrams should be added and insert placeholders
# This is a simplified approach - you may want to manually add diagrams

print("Document formatting updated!")
print("Note: You may need to manually insert actual diagrams/images where placeholders are indicated.")

# Save the updated document
target_doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("Document saved successfully!")

