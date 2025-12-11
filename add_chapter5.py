from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Add page break before Chapter 5
doc.add_page_break()

# CHAPTER - 5 INPUT / OUTPUT DESIGN
chapter5_heading = doc.add_heading('CHAPTER - 5', level=1)
chapter5_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
chapter5_title = doc.add_heading('INPUT / OUTPUT DESIGN', level=1)
chapter5_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# INPUT DESIGN
doc.add_heading('5.1 INPUT DESIGN', level=2)

input_design_text = """Input design focuses on creating user-friendly forms and data entry interfaces that facilitate accurate and efficient data collection. The design principles ensure data integrity, user convenience, and error prevention.

**INPUT DESIGN PRINCIPLES:**

1. **Simplicity**: Forms should be easy to understand and complete
2. **Validation**: Real-time and server-side validation to prevent errors
3. **Feedback**: Clear error messages and success indicators
4. **Accessibility**: Keyboard navigation and screen reader support
5. **Consistency**: Uniform form patterns across the application

**REGISTRATION FORM INPUT DESIGN:**

**Form Name:** User Registration Form
**Location:** /register
**Purpose:** New user account creation

**Input Fields:**

1. **Name Field**
   - **Type**: Text input
   - **Label**: "Name"
   - **Placeholder**: None
   - **Required**: Yes
   - **Validation**: 
     - Minimum 2 characters
     - Maximum 100 characters
     - Only letters and spaces allowed
   - **Error Messages**: 
     - "Name must be at least 2 characters"
     - "Name is required"
   - **Design**: Full-width input with focus ring

2. **Email Field**
   - **Type**: Email input
   - **Label**: "Email"
   - **Placeholder**: None
   - **Required**: Yes
   - **Validation**: 
     - Valid email format
     - Unique email (checked on server)
   - **Error Messages**: 
     - "Please enter a valid email address"
     - "Email already exists"
   - **Design**: Full-width input with email icon

3. **Password Field**
   - **Type**: Password input
   - **Label**: "Password"
   - **Placeholder**: None
   - **Required**: Yes
   - **Validation**: 
     - Minimum 8 characters
     - Must contain uppercase letter
     - Must contain lowercase letter
     - Must contain number
   - **Error Messages**: 
     - "Password must be at least 8 characters"
     - "Password must contain uppercase, lowercase, and number"
   - **Design**: Full-width input with show/hide toggle (optional)

4. **Role Selection**
   - **Type**: Select dropdown
   - **Label**: "Role"
   - **Options**: "Student", "Recruiter"
   - **Required**: Yes
   - **Default**: "Student"
   - **Design**: Dropdown with clear options

5. **College Field** (Conditional - Student only)
   - **Type**: Text input
   - **Label**: "College"
   - **Placeholder**: "Enter your college name"
   - **Required**: Yes (for students)
   - **Validation**: 
     - Minimum 2 characters
     - Maximum 200 characters
   - **Error Messages**: 
     - "College name is required"
     - "College name must be at least 2 characters"
   - **Design**: Appears only when role is "Student"

6. **Company Name Field** (Conditional - Recruiter only)
   - **Type**: Text input
   - **Label**: "Company Name"
   - **Placeholder**: "Enter company name (optional)"
   - **Required**: No
   - **Validation**: 
     - Maximum 200 characters (if provided)
   - **Design**: Appears only when role is "Recruiter"

**Form Layout:**
- Vertical form layout
- Fields stacked vertically
- Submit button at bottom
- Link to login page below form
- Error messages displayed below respective fields
- Success message via toast notification

**LOGIN FORM INPUT DESIGN:**

**Form Name:** User Login Form
**Location:** /login
**Purpose:** User authentication

**Input Fields:**

1. **Email Field**
   - **Type**: Email input
   - **Label**: "Email"
   - **Placeholder**: "your.email@example.com"
   - **Required**: Yes
   - **Validation**: Valid email format
   - **Error Messages**: "Please enter a valid email address"

2. **Password Field**
   - **Type**: Password input
   - **Label**: "Password"
   - **Placeholder**: "Enter your password"
   - **Required**: Yes
   - **Validation**: Non-empty
   - **Error Messages**: "Password is required"

**Form Layout:**
- Simple two-field form
- Submit button
- "Forgot Password" link (if implemented)
- Link to registration page
- Error messages displayed below form

**JOB POSTING FORM INPUT DESIGN:**

**Form Name:** Job Posting Form
**Location:** /post-job
**Purpose:** Create new job posting (Recruiters only)

**Input Fields:**

1. **Job Title Field**
   - **Type**: Text input
   - **Label**: "Title"
   - **Placeholder**: "e.g., Full Stack Developer"
   - **Required**: Yes
   - **Validation**: 
     - Minimum 3 characters
     - Maximum 200 characters
   - **Error Messages**: "Job title is required"

2. **Job Description Field**
   - **Type**: Textarea
   - **Label**: "Description"
   - **Placeholder**: "Enter detailed job description..."
   - **Required**: Yes
   - **Rows**: 6
   - **Validation**: 
     - Minimum 20 characters
     - Maximum 5000 characters
   - **Error Messages**: "Description must be at least 20 characters"

3. **Job Type Field**
   - **Type**: Select dropdown
   - **Label**: "Type"
   - **Options**: "Internship", "Full-Time"
   - **Required**: Yes
   - **Default**: "Internship"
   - **Design**: Dropdown selection

4. **Salary/Stipend Field**
   - **Type**: Text input
   - **Label**: "Salary/Stipend"
   - **Placeholder**: "e.g., ₹50,000 - ₹80,000 or ₹15,000/month"
   - **Required**: Yes
   - **Validation**: Non-empty
   - **Error Messages**: "Salary/Stipend is required"

5. **Location Field**
   - **Type**: Text input
   - **Label**: "Location"
   - **Placeholder**: "Ahmedabad"
   - **Required**: Yes
   - **Default**: "Ahmedabad"
   - **Validation**: Non-empty

6. **Required Skills Field**
   - **Type**: Text input
   - **Label**: "Required Skills (comma separated)"
   - **Placeholder**: "React, Node.js, MongoDB"
   - **Required**: Yes
   - **Validation**: Non-empty
   - **Processing**: Split by comma, trim, filter empty
   - **Error Messages**: "At least one skill is required"

**Form Layout:**
- Multi-field form
- Textarea for description
- Submit and Cancel buttons
- Company validation check (must have company profile)

**INTERVIEW EXPERIENCE FORM INPUT DESIGN:**

**Form Name:** Interview Experience Submission Form
**Location:** /submit-interview
**Purpose:** Submit interview experience (Students only)

**Input Fields:**

1. **Company Selection**
   - **Type**: Select dropdown
   - **Label**: "Company"
   - **Options**: Populated from companies list
   - **Required**: Yes
   - **Validation**: Must select valid company
   - **Error Messages**: "Please select a company"

2. **Interview Date Field**
   - **Type**: Date input
   - **Label**: "Interview Date"
   - **Placeholder**: None
   - **Required**: No
   - **Default**: Current date
   - **Validation**: Valid date format

3. **Interview Rounds Section** (Dynamic)
   - **Type**: Dynamic form section
   - **Label**: "Interview Rounds"
   - **Required**: At least one round
   - **Fields per Round**:
     - **Round Type**: Text input (e.g., "Technical", "HR", "Coding Test")
     - **Questions Asked**: Textarea (semicolon-separated)
     - **Difficulty Rating**: Number input (1-5, slider or select)
   - **Actions**: Add Round button, Remove Round button
   - **Validation**: 
     - Round type required
     - At least one question required
     - Difficulty rating required (1-5)

4. **Overall Rating Field**
   - **Type**: Number input (slider or select)
   - **Label**: "Overall Rating"
   - **Range**: 1-5
   - **Required**: Yes
   - **Default**: 3
   - **Design**: Star rating or slider

5. **Anonymous Checkbox**
   - **Type**: Checkbox
   - **Label**: "Post anonymously"
   - **Required**: No
   - **Default**: Checked (true)
   - **Design**: Checkbox with label

**Form Layout:**
- Multi-section form
- Dynamic rounds section with add/remove functionality
- Submit button
- Complex validation for nested data

**PROFILE UPDATE FORM INPUT DESIGN:**

**Form Name:** User Profile Update Form
**Location:** /profile
**Purpose:** Update user profile information

**Input Fields:**

1. **Name Field**
   - **Type**: Text input
   - **Label**: "Name"
   - **Required**: Yes
   - **Validation**: Minimum 2 characters
   - **Design**: Editable field

2. **Email Field**
   - **Type**: Email input
   - **Label**: "Email"
   - **Required**: Yes
   - **Read-only**: Yes (cannot be changed)
   - **Design**: Disabled/grayed out field

3. **College Field** (Students only)
   - **Type**: Text input
   - **Label**: "College"
   - **Required**: No
   - **Validation**: Minimum 2 characters (if provided)
   - **Design**: Editable field

4. **Company Field** (Recruiters only)
   - **Type**: Text display
   - **Label**: "Company"
   - **Read-only**: Yes
   - **Design**: Display only, link to company page

5. **Avatar Upload**
   - **Type**: File input
   - **Label**: "Profile Picture"
   - **Accept**: image/* (jpg, png, gif, webp)
   - **Required**: No
   - **Validation**: 
     - File size: Max 5MB
     - File type: Image only
   - **Design**: 
     - Current avatar preview
     - Upload button
     - Image preview after selection
   - **Error Messages**: 
     - "File size must be less than 5MB"
     - "Please select an image file"

**Form Layout:**
- Profile information section
- Avatar upload section (separate)
- Update button
- Separate API call for avatar upload

**COMPANY CREATION FORM INPUT DESIGN:**

**Form Name:** Company Profile Creation Form
**Location:** /create-company
**Purpose:** Create company profile (Recruiters only)

**Input Fields:**

1. **Company Name Field**
   - **Type**: Text input
   - **Label**: "Company Name"
   - **Required**: Yes
   - **Validation**: 
     - Minimum 2 characters
     - Unique company name
   - **Error Messages**: "Company name already exists"

2. **Description Field**
   - **Type**: Textarea
   - **Label**: "Description"
   - **Required**: Yes
   - **Rows**: 6
   - **Validation**: Minimum 20 characters

3. **Website Field**
   - **Type**: URL input
   - **Label**: "Website"
   - **Required**: No
   - **Placeholder**: "https://www.example.com"
   - **Validation**: Valid URL format (if provided)

4. **Address Field**
   - **Type**: Textarea
   - **Label**: "Address"
   - **Required**: Yes
   - **Rows**: 3
   - **Validation**: Minimum 5 characters

5. **Tech Stack Field**
   - **Type**: Text input
   - **Label**: "Tech Stack (comma separated)"
   - **Placeholder**: "React, Node.js, MongoDB, Python"
   - **Required**: No
   - **Processing**: Split by comma, trim, filter empty

6. **Logo Upload**
   - **Type**: File input
   - **Label**: "Company Logo"
   - **Accept**: image/*
   - **Required**: No
   - **Validation**: 
     - File size: Max 5MB
     - File type: Image only
   - **Design**: 
     - Upload button
     - Image preview after selection

**Form Layout:**
- Multi-field form
- File upload section
- Submit button

**INPUT VALIDATION STRATEGY:**

**Client-Side Validation:**
- Real-time field validation
- Format checking (email, URL)
- Length validation
- Required field checking
- Immediate feedback

**Server-Side Validation:**
- Duplicate checking (email, company name)
- Business rule validation
- Data integrity checks
- Security validation
- Comprehensive error messages

**Error Display:**
- Inline error messages below fields
- Toast notifications for form submission errors
- Color-coded error indicators (red borders)
- Success messages via toast notifications

**INPUT DESIGN BEST PRACTICES:**

1. **Clear Labels**: Descriptive labels for all fields
2. **Placeholders**: Helpful placeholder text where appropriate
3. **Required Indicators**: Asterisk (*) for required fields
4. **Field Grouping**: Logical grouping of related fields
5. **Progressive Disclosure**: Show/hide fields based on selections
6. **Auto-fill Support**: Appropriate input types for browser auto-fill
7. **Keyboard Navigation**: Tab order and Enter key submission
8. **Loading States**: Disable form during submission
9. **Confirmation**: Success messages after submission
10. **Accessibility**: ARIA labels and screen reader support

This comprehensive input design ensures accurate data collection, prevents errors, and provides a smooth user experience throughout the Ahmedabad Career Hub application."""

doc.add_paragraph(input_design_text)

# OUTPUT DESIGN
doc.add_heading('5.2 OUTPUT DESIGN', level=2)

output_design_text = """Output design focuses on presenting information to users in a clear, organized, and visually appealing manner. The design ensures that users can easily understand and interact with the displayed data.

**OUTPUT DESIGN PRINCIPLES:**

1. **Clarity**: Information should be easy to read and understand
2. **Organization**: Logical grouping and hierarchy of information
3. **Visual Hierarchy**: Important information stands out
4. **Consistency**: Uniform display patterns across the application
5. **Responsiveness**: Adapts to different screen sizes

**JOB LISTING OUTPUT DESIGN:**

**Output Name:** Job Listings Display
**Location:** /jobs
**Purpose:** Display available job postings

**Output Components:**

**Job Card Design:**
- **Layout**: Card-based grid layout
- **Card Elements**:
  - Job Title (Large, bold, blue)
  - Company Name (Medium, link to company)
  - Job Type Badge (Internship/Full-time, color-coded)
  - Location Icon and Text
  - Required Skills (Tags/chips)
  - Salary/Stipend (Highlighted)
  - Posted Date (Small, gray)
  - "View Details" button/link
- **Card Styling**:
  - White background
  - Border and shadow
  - Hover effect (lift and shadow increase)
  - Responsive grid (1 column mobile, 2-3 columns desktop)

**List View Design:**
- Grid layout with consistent spacing
- Pagination controls at bottom
- Search and filter bar at top
- Empty state message when no jobs found
- Loading skeleton while fetching

**Output Format:**
```
┌─────────────────────────────────────┐
│ Full Stack Developer               │
│ Tech Solutions Inc.                │
│ [Internship] [Ahmedabad]           │
│ Skills: React, Node.js, MongoDB    │
│ ₹50,000 - ₹80,000                  │
│ Posted: 2 days ago                 │
│ [View Details →]                    │
└─────────────────────────────────────┘
```

**JOB DETAIL OUTPUT DESIGN:**

**Output Name:** Job Detail Page
**Location:** /jobs/:id
**Purpose:** Display complete job information

**Output Components:**

**Header Section:**
- Job Title (Extra large, bold)
- Company Name with logo (Large, link to company)
- Job Type Badge
- Location
- Posted Date

**Main Content Section:**
- **Job Description**: 
  - Full text description
  - Formatted paragraphs
  - Readable font size and line height
- **Job Details Box**:
  - Job Type
  - Location
  - Salary/Stipend
  - Required Skills (List or tags)
- **Company Information Card**:
  - Company name and logo
  - Company description (preview)
  - Link to full company page

**Action Section:**
- Apply Button (for students, if not applied)
- Application Status (if already applied)
- Status badge with color coding:
  - Applied: Blue
  - Reviewed: Yellow
  - Shortlisted: Green
  - Rejected: Red
  - Hired: Green

**Output Format:**
```
┌─────────────────────────────────────┐
│ Full Stack Developer                │
│ Tech Solutions Inc.                  │
│ [Full-Time] Ahmedabad               │
├─────────────────────────────────────┤
│ Job Description:                    │
│ [Full description text...]           │
├─────────────────────────────────────┤
│ Details:                            │
│ • Type: Full-Time                   │
│ • Location: Ahmedabad               │
│ • Salary: ₹50,000 - ₹80,000        │
│ • Skills: React, Node.js, MongoDB   │
├─────────────────────────────────────┤
│ [Apply Now] or [Status: Applied]    │
└─────────────────────────────────────┘
```

**COMPANY LISTING OUTPUT DESIGN:**

**Output Name:** Company Listings Display
**Location:** /companies
**Purpose:** Display available companies

**Output Components:**

**Company Card Design:**
- **Layout**: Card-based grid layout
- **Card Elements**:
  - Company Logo (Circular or square)
  - Company Name (Large, bold)
  - Description Preview (Truncated, 2-3 lines)
  - Tech Stack Tags (Chips/badges)
  - Location (if available)
  - "View Details" button/link
- **Card Styling**:
  - White background
  - Border and shadow
  - Hover effect
  - Responsive grid

**Output Format:**
```
┌─────────────────────────────────────┐
│ [Logo]                              │
│ Tech Solutions Inc.                 │
│ Leading software development...     │
│ [React] [Node.js] [MongoDB]          │
│ Ahmedabad                           │
│ [View Details →]                    │
└─────────────────────────────────────┘
```

**COMPANY DETAIL OUTPUT DESIGN:**

**Output Name:** Company Detail Page
**Location:** /companies/:id
**Purpose:** Display complete company information

**Output Components:**

**Header Section:**
- Company Logo (Large)
- Company Name (Extra large, bold)
- Website Link (if available)
- Location/Address

**Main Content Section:**
- **Company Description**: 
  - Full text description
  - Formatted paragraphs
- **Company Details Box**:
  - Website
  - Address
  - Tech Stack (Tags/chips)
- **Available Jobs Section**:
  - List of jobs posted by company
  - Job cards (compact)
  - Link to job details
- **Interview Experiences Section**:
  - List of interview experiences
  - Rating display
  - Link to view details

**Output Format:**
```
┌─────────────────────────────────────┐
│ [Large Logo]                         │
│ Tech Solutions Inc.                 │
│ www.techsolutions.com                │
├─────────────────────────────────────┤
│ Company Description:                │
│ [Full description...]                │
├─────────────────────────────────────┤
│ Tech Stack:                          │
│ [React] [Node.js] [MongoDB] [Python] │
├─────────────────────────────────────┤
│ Available Jobs (3):                  │
│ • Full Stack Developer              │
│ • Frontend Developer                 │
│ • Backend Developer                  │
├─────────────────────────────────────┤
│ Interview Experiences (5):          │
│ [Rating: 4.2/5]                      │
└─────────────────────────────────────┘
```

**APPLICATION LIST OUTPUT DESIGN:**

**Output Name:** Application History Display
**Location:** /applications
**Purpose:** Display student's job applications

**Output Components:**

**Application Card Design:**
- **Layout**: List or card layout
- **Card Elements**:
  - Job Title (Large, bold, link to job)
  - Company Name (Medium, link to company)
  - Application Status Badge (Color-coded):
    - Applied: Blue
    - Reviewed: Yellow
    - Shortlisted: Green
    - Rejected: Red
    - Hired: Green (with checkmark)
  - Applied Date (Small, gray)
  - Status Update Date (if status changed)
- **Card Styling**:
  - White background
  - Border
  - Status badge prominently displayed
  - Hover effect

**List View Design:**
- Vertical list layout
- Filter by status option
- Pagination controls
- Empty state when no applications
- Loading state while fetching

**Output Format:**
```
┌─────────────────────────────────────┐
│ Full Stack Developer                │
│ Tech Solutions Inc.                 │
│ [Status: Shortlisted] ✓             │
│ Applied: Jan 15, 2024               │
│ Updated: Jan 20, 2024               │
│ [View Job Details →]                │
└─────────────────────────────────────┘
```

**DASHBOARD OUTPUT DESIGN:**

**Output Name:** Student Dashboard
**Location:** /dashboard
**Purpose:** Overview of student activities

**Output Components:**

**Welcome Section:**
- Personalized greeting: "Welcome, [Name]!"
- User avatar (if available)

**Recent Applications Section:**
- Section title: "Recent Applications"
- Application cards (shows 3 most recent)
- "View All Applications" button
- Empty state if no applications

**Quick Actions Section:**
- Action buttons:
  - "View All Applications"
  - "Submit Interview Experience"
  - "Browse Jobs"
  - "Browse Companies"
- Icon-based buttons
- Hover effects

**Statistics Section (Optional):**
- Total applications
- Applications by status
- Charts/graphs (if implemented)

**Output Format:**
```
┌─────────────────────────────────────┐
│ Welcome, John Doe!                  │
├─────────────────────────────────────┤
│ Recent Applications:                │
│ [Application Card 1]                │
│ [Application Card 2]                │
│ [Application Card 3]                │
│ [View All →]                        │
├─────────────────────────────────────┤
│ Quick Actions:                       │
│ [Applications] [Interview] [Jobs]   │
└─────────────────────────────────────┘
```

**RECRUITER DASHBOARD OUTPUT DESIGN:**

**Output Name:** Recruiter Dashboard
**Location:** /recruiter-dashboard
**Purpose:** Overview of recruiter activities

**Output Components:**

**Welcome Section:**
- Personalized greeting
- Company information

**Posted Jobs Section:**
- Section title: "Your Posted Jobs"
- Job list/table:
  - Job Title
  - Job Type
  - Applications Count
  - Posted Date
  - Actions (Edit, Delete, View Applications)
- "Post New Job" button
- Empty state if no jobs

**Quick Actions Section:**
- Action buttons:
  - "Post New Job"
  - "Manage Applications"
  - "View My Company"
  - "Create Company" (if no company)

**Statistics Section:**
- Total jobs posted
- Total applications received
- Applications by status

**Output Format:**
```
┌─────────────────────────────────────┐
│ Welcome, Recruiter!                 │
│ Tech Solutions Inc.                 │
├─────────────────────────────────────┤
│ Your Posted Jobs:                   │
│ • Full Stack Developer (5 apps)    │
│ • Frontend Developer (3 apps)       │
│ [Post New Job]                      │
├─────────────────────────────────────┤
│ Quick Actions:                       │
│ [Post Job] [Manage] [Company]      │
└─────────────────────────────────────┘
```

**INTERVIEW EXPERIENCE OUTPUT DESIGN:**

**Output Name:** Interview Experience Display
**Location:** Company detail page, Interview listings
**Purpose:** Display interview experiences

**Output Components:**

**Interview Card Design:**
- **Layout**: Card layout
- **Card Elements**:
  - Company Name (if not anonymous)
  - Interview Date
  - Overall Rating (Stars: 1-5)
  - Rounds Summary:
    - Round Type
    - Difficulty Rating
    - Number of questions
  - "View Details" link
- **Card Styling**:
  - White background
  - Border
  - Rating prominently displayed

**Interview Detail View:**
- Company information
- Interview date
- Overall rating (visual stars)
- Detailed rounds:
  - Round type
  - Questions asked (list)
  - Difficulty rating
- Anonymous indicator (if applicable)

**Output Format:**
```
┌─────────────────────────────────────┐
│ Interview Experience                │
│ Tech Solutions Inc.                  │
│ Date: Jan 15, 2024                  │
│ Rating: ★★★★☆ (4/5)                │
├─────────────────────────────────────┤
│ Rounds:                              │
│ 1. Technical (Difficulty: 4/5)      │
│    Questions:                        │
│    • What is React?                  │
│    • Explain closures                │
│ 2. HR (Difficulty: 3/5)              │
│    Questions:                        │
│    • Tell us about yourself         │
└─────────────────────────────────────┘
```

**OUTPUT DESIGN FEATURES:**

**Visual Elements:**
- Color-coded status badges
- Icons for visual clarity
- Rating displays (stars, numbers)
- Progress indicators
- Loading states (skeletons, spinners)

**Typography:**
- Clear hierarchy (headings, subheadings, body)
- Readable font sizes
- Appropriate line heights
- Color contrast for readability

**Layout:**
- Responsive grid systems
- Consistent spacing
- Card-based layouts
- List and table views
- Mobile-first design

**Interactive Elements:**
- Hover effects
- Clickable cards/links
- Action buttons
- Pagination controls
- Filter and search interfaces

**Empty States:**
- Friendly messages when no data
- Call-to-action buttons
- Illustrations or icons

**Error States:**
- Clear error messages
- Retry options
- Helpful guidance

**Loading States:**
- Skeleton loaders
- Spinner animations
- Progress indicators
- Disabled states during loading

**OUTPUT DESIGN BEST PRACTICES:**

1. **Information Hierarchy**: Most important information first
2. **Visual Clarity**: Clear separation between sections
3. **Consistent Styling**: Uniform design patterns
4. **Responsive Design**: Works on all screen sizes
5. **Accessibility**: Screen reader support, keyboard navigation
6. **Performance**: Fast loading, optimized images
7. **User Feedback**: Loading, success, and error states
8. **Pagination**: For large datasets
9. **Search and Filter**: Easy data discovery
10. **Print-Friendly**: Consider print styles (if applicable)

This comprehensive output design ensures that information is presented clearly, attractively, and in a user-friendly manner throughout the Ahmedabad Career Hub application."""

doc.add_paragraph(output_design_text)

# Save the updated document
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("Chapter 5 has been added successfully to the document!")

