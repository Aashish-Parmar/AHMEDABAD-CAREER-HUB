from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Inserting diagram placeholders at appropriate sections...")

doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Function to create diagram with proper formatting (matching reference style)
def create_diagram(doc, figure_num, title):
    """Create a diagram placeholder matching reference document style"""
    
    # Add spacing before
    doc.add_paragraph()
    
    # Figure title (centered, bold, Times New Roman 11pt)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    
    title_text = f"Figure {figure_num}: {title}"
    title_run = title_para.add_run(title_text)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(11)
    title_run.bold = True
    
    # Small spacing
    doc.add_paragraph()
    
    # Create bordered table for diagram
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table.columns[0].width = Inches(5.5)
    cell = table.rows[0].cells[0]
    cell.height = Inches(3)
    
    # Add borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'a:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)
    
    # Placeholder text
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    placeholder_run = cell_para.add_run("[Insert Diagram Here]")
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Spacing after
    doc.add_paragraph()
    doc.add_paragraph()

# Apply consistent formatting
print("Applying consistent font formatting...")
for para in doc.paragraphs:
    if para.runs:
        for run in para.runs:
            if not run.font.name or run.font.name in ['Calibri', 'Arial', None]:
                run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

# Find specific sections and add diagrams
print("Locating sections for diagram insertion...")

# Track figure numbers
figure_counter = {
    '3': 1,  # Chapter 3 figures
    '4': 1,  # Chapter 4 figures
}

# Find sections and insert diagrams
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Chapter 3.3.1 - WBS
    if "3.3.1 WORK BREAKDOWN STRUCTURE" in text.upper():
        # Find end of this section (before next heading)
        for j in range(i+1, min(i+50, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.style.name.startswith('Heading') and j > i+5:
                # Insert diagram before next heading
                create_diagram(doc.paragraphs[j], f"3.{figure_counter['3']}", 
                             "Work Breakdown Structure Diagram")
                figure_counter['3'] += 1
                break
    
    # Chapter 3.3.2 - Gantt Chart
    if "3.3.2 GANTT CHART" in text.upper():
        for j in range(i+1, min(i+50, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.style.name.startswith('Heading') and j > i+5:
                create_diagram(doc.paragraphs[j], f"3.{figure_counter['3']}", 
                             "Gantt Chart - Project Timeline")
                figure_counter['3'] += 1
                break
    
    # Chapter 4.1 - UML/DFD
    if "4.1 UML" in text.upper() or "4.1 UNIFIED MODELING LANGUAGE" in text.upper():
        # Add multiple diagrams for this section
        for j in range(i+1, min(i+100, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if "4.2 SYSTEM FLOW" in next_para.text.upper() and j > i+10:
                # Insert diagrams before 4.2
                diagrams = [
                    "Entity Relationship Diagram (ERD)",
                    "Data Flow Diagram - Level 0 (Context Diagram)",
                    "Data Flow Diagram - Level 1",
                    "Use Case Diagram",
                ]
                insert_pos = j
                for idx, diag_title in enumerate(diagrams):
                    create_diagram(doc.paragraphs[insert_pos], 
                                 f"4.{figure_counter['4']}", diag_title)
                    figure_counter['4'] += 1
                    insert_pos += 15  # Space between diagrams
                break
    
    # Chapter 4.2 - System Flow
    if "4.2 SYSTEM FLOW" in text.upper():
        for j in range(i+1, min(i+100, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if "4.3 DATA DICTIONARY" in next_para.text.upper() and j > i+10:
                # Insert flow diagrams
                flow_diagrams = [
                    "User Registration Flow Diagram",
                    "User Login Flow Diagram",
                    "Job Browsing and Application Flow Diagram",
                ]
                insert_pos = j
                for diag_title in flow_diagrams:
                    create_diagram(doc.paragraphs[insert_pos], 
                                 f"4.{figure_counter['4']}", diag_title)
                    figure_counter['4'] += 1
                    insert_pos += 15
                break

# Save
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print(f"\nDocument updated! Added diagram placeholders.")
print(f"Total figures added: Chapter 3: {figure_counter['3']-1}, Chapter 4: {figure_counter['4']-1}")
print("\nNote: Replace placeholder boxes with actual diagrams/images.")

