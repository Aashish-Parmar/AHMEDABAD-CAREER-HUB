from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Add Title Page
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('AHMEDABAD CAREER HUB')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run('A Full-Stack Job Portal Application')
subtitle_run.font.size = Pt(18)
subtitle_run.font.bold = True

doc.add_paragraph()  # Add space

author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_para.add_run('Project Documentation')
author_run.font.size = Pt(14)

doc.add_page_break()

# CHAPTER - 1 INTRODUCTION
chapter1_heading = doc.add_heading('CHAPTER - 1', level=1)
chapter1_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
chapter1_title = doc.add_heading('INTRODUCTION', level=1)
chapter1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1.1 ORGANIZATION PROFILE
doc.add_heading('1.1 ORGANIZATION PROFILE', level=2)

org_text = """Ahmedabad Career Hub is a comprehensive job portal platform designed specifically to bridge the gap between students seeking career opportunities and recruiters looking for talented candidates in the Ahmedabad region. The platform serves as a centralized ecosystem where educational institutions, students, and companies converge to facilitate seamless job placement and career development.

The organization focuses on creating a user-friendly, efficient, and transparent job market that benefits all stakeholders. By leveraging modern web technologies and best practices in software development, Ahmedabad Career Hub aims to revolutionize how students discover opportunities and how companies find the right talent.

The platform operates with a mission to empower students with the tools and information they need to make informed career decisions, while providing recruiters with an efficient system to manage job postings and candidate applications. This dual-sided marketplace approach ensures value creation for both parties involved in the recruitment process."""

doc.add_paragraph(org_text)

# 1.2 SYSTEM DETAILS
doc.add_heading('1.2 SYSTEM DETAILS', level=2)

# 1.2.1 EXISTING SYSTEM
doc.add_heading('1.2.1 EXISTING SYSTEM', level=3)

existing_text = """The traditional job search and recruitment process in Ahmedabad, like in many other cities, faces several challenges:

1. **Fragmented Job Listings**: Job opportunities are scattered across multiple platforms, making it difficult for students to find relevant positions in one place.

2. **Limited Information Sharing**: Students lack access to real interview experiences and insights from peers who have gone through the recruitment process at specific companies.

3. **Manual Application Process**: The application process is often tedious, requiring students to repeatedly fill out forms and submit documents for each position.

4. **Lack of Centralized Company Information**: Students struggle to find comprehensive information about companies, their culture, tech stack, and work environment.

5. **Inefficient Recruiter Tools**: Recruiters face challenges in managing multiple job postings, tracking applications, and maintaining company profiles across different platforms.

6. **No Direct Student-Recruiter Connection**: There is no dedicated platform that facilitates direct communication and interaction between students and recruiters in the Ahmedabad region.

7. **Limited Career Guidance**: Students often lack access to structured career guidance, interview preparation resources, and industry insights specific to their region.

These limitations result in a time-consuming, inefficient, and often frustrating experience for both job seekers and recruiters, leading to missed opportunities and suboptimal hiring decisions."""

doc.add_paragraph(existing_text)

# 1.2.2 PROPOSED SYSTEM
doc.add_heading('1.2.2 PROPOSED SYSTEM', level=3)

proposed_text = """Ahmedabad Career Hub is a modern, full-stack web application designed to address all the limitations of the existing system. The proposed system offers:

**1. Unified Job Portal**
   - Centralized platform where all job opportunities in Ahmedabad are listed
   - Advanced search and filtering capabilities based on job type, skills, location, and company
   - Real-time job updates and notifications

**2. Student-Centric Features**
   - Easy job browsing and application submission
   - Application tracking system to monitor status of submitted applications
   - Interview experience sharing platform where students can post and view real interview experiences
   - Comprehensive company profiles with detailed information about tech stack, culture, and work environment
   - User profile management with avatar upload functionality

**3. Recruiter Management Tools**
   - Company profile creation and management
   - Job posting interface for internships and full-time positions
   - Application management dashboard to view and track candidate applications
   - Company-specific interview experience repository

**4. Technical Architecture**
   - **Frontend**: Built with React 19, Vite, and TailwindCSS for a modern, responsive user interface
   - **Backend**: Node.js with Express 5 framework for robust API development
   - **Database**: MongoDB with Mongoose ODM for flexible data management
   - **Authentication**: JWT-based secure authentication system
   - **File Management**: Multer for handling file uploads (avatars, documents)

**5. User Experience Enhancements**
   - Role-based access control (Student/Recruiter)
   - Responsive design for desktop and mobile devices
   - Intuitive navigation and user-friendly interface
   - Real-time feedback and error handling
   - Secure data handling and privacy protection

**6. Key Differentiators**
   - Region-specific focus on Ahmedabad job market
   - Interview experience sharing feature unique to the platform
   - Seamless integration between job search, application, and interview preparation
   - Community-driven insights and experiences

The proposed system transforms the traditional job search process into a streamlined, efficient, and user-friendly experience that benefits both students and recruiters."""

doc.add_paragraph(proposed_text)

# 1.3 SCOPE OF SYSTEM
doc.add_heading('1.3 SCOPE OF SYSTEM', level=2)

scope_text = """The scope of Ahmedabad Career Hub encompasses the following areas:

**1. User Management**
   - User registration and authentication for students and recruiters
   - Profile management including personal information, college details, and avatar upload
   - Role-based access control and authorization

**2. Job Management**
   - Job posting creation and management by recruiters
   - Job browsing, searching, and filtering by students
   - Support for both internship and full-time positions
   - Job details including title, description, salary/stipend, location, and required skills

**3. Company Management**
   - Company profile creation and management
   - Company information including name, description, website, logo, address, and tech stack
   - Company listing and detailed company pages

**4. Application Management**
   - Job application submission by students
   - Application status tracking (applied, reviewed, shortlisted, rejected, hired)
   - Application history and management for both students and recruiters

**5. Interview Experience Platform**
   - Interview experience submission by students
   - Detailed interview information including rounds, questions asked, and difficulty ratings
   - Company-specific interview experience repository
   - Anonymous posting option for privacy

**6. Search and Discovery**
   - Advanced job search with multiple filter options
   - Company search and browsing
   - Interview experience search by company

**7. Technical Scope**
   - RESTful API development for all system functionalities
   - Secure authentication and authorization
   - File upload and management
   - Database design and optimization
   - Responsive frontend development
   - Error handling and validation

**Geographical Scope**: Initially focused on Ahmedabad region, with potential for expansion to other cities.

**User Scope**: 
   - Students from various colleges and universities in Ahmedabad
   - Recruiters and HR professionals from companies operating in Ahmedabad
   - Companies ranging from startups to established enterprises

**Functional Scope**: The system covers the complete job search and recruitment lifecycle from job posting to application submission and interview experience sharing."""

doc.add_paragraph(scope_text)

# 1.4 OBJECTIVES
doc.add_heading('1.4 OBJECTIVES', level=2)

objectives_text = """The primary and secondary objectives of Ahmedabad Career Hub are as follows:

**PRIMARY OBJECTIVES:**

1. **To Create a Centralized Job Portal**
   - Develop a single platform where all job opportunities in Ahmedabad are aggregated
   - Eliminate the need for students to visit multiple websites for job search
   - Provide recruiters with a dedicated platform for posting opportunities

2. **To Streamline the Application Process**
   - Simplify job application submission for students
   - Enable one-click application process with stored user information
   - Provide real-time application status tracking

3. **To Facilitate Information Sharing**
   - Create a platform for students to share real interview experiences
   - Enable access to authentic interview questions and company insights
   - Build a knowledge repository for interview preparation

4. **To Connect Students and Recruiters**
   - Bridge the gap between job seekers and employers
   - Facilitate direct communication and interaction
   - Create a transparent recruitment ecosystem

**SECONDARY OBJECTIVES:**

1. **To Provide Comprehensive Company Information**
   - Offer detailed company profiles with tech stack, culture, and work environment
   - Enable students to make informed decisions about potential employers
   - Help recruiters showcase their company effectively

2. **To Enhance User Experience**
   - Develop an intuitive and user-friendly interface
   - Ensure responsive design for all devices
   - Implement fast and efficient system performance

3. **To Ensure Data Security and Privacy**
   - Implement secure authentication mechanisms
   - Protect user data and privacy
   - Ensure compliance with data protection standards

4. **To Support Career Development**
   - Provide tools for students to track their job search progress
   - Enable recruiters to efficiently manage their hiring process
   - Create opportunities for skill development and career growth

5. **To Build a Regional Job Market Ecosystem**
   - Focus on Ahmedabad-specific opportunities
   - Support local companies and educational institutions
   - Contribute to regional economic development

**SUCCESS CRITERIA:**

- Successful user registration and authentication for both students and recruiters
- Efficient job posting and management functionality
- Seamless application submission and tracking
- Active interview experience sharing community
- High user satisfaction and engagement
- System reliability and performance
- Secure data handling and privacy protection"""

doc.add_paragraph(objectives_text)

# Save the document
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("Chapter 1 documentation has been created successfully as 'Ahmedabad_Career_Hub_Chapter1.docx'")

