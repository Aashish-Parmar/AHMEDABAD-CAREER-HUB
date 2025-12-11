from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Add page break before Chapter 4
doc.add_page_break()

# CHAPTER - 4 SYSTEM ANALYSIS AND DESIGN
chapter4_heading = doc.add_heading('CHAPTER - 4', level=1)
chapter4_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
chapter4_title = doc.add_heading('SYSTEM ANALYSIS AND DESIGN', level=1)
chapter4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 4.1 UML (UNIFIED MODELING LANGUAGE) / DFD
doc.add_heading('4.1 UML (UNIFIED MODELING LANGUAGE) / DFD', level=2)

uml_dfd_text = """This section describes the system design using UML (Unified Modeling Language) diagrams and Data Flow Diagrams (DFD) to represent the system architecture, data flow, and relationships.

**UML CLASS DIAGRAM:**

The class diagram represents the data models and their relationships in the Ahmedabad Career Hub system.

**ENTITIES AND RELATIONSHIPS:**

**1. User Entity**
- Attributes: _id, name, email, password, role, college, companyName, company (ObjectId), avatarUrl, createdAt, updatedAt
- Relationships:
  - One-to-Many with Application (a user can have multiple applications)
  - One-to-Many with Interview (a user can submit multiple interview experiences)
  - Many-to-One with Company (a recruiter belongs to one company)

**2. Company Entity**
- Attributes: _id, companyName, description, website, logoUrl, address, techStack[], createdAt, updatedAt
- Relationships:
  - One-to-Many with User (a company can have multiple recruiters)
  - One-to-Many with Job (a company can post multiple jobs)
  - One-to-Many with Interview (a company can have multiple interview experiences)

**3. Job Entity**
- Attributes: _id, title, description, jobType, salaryStipend, location, requiredSkills[], company (ObjectId), createdAt, updatedAt
- Relationships:
  - Many-to-One with Company (a job belongs to one company)
  - One-to-Many with Application (a job can have multiple applications)

**4. Application Entity**
- Attributes: _id, user (ObjectId), job (ObjectId), appliedAt, status
- Relationships:
  - Many-to-One with User (an application belongs to one user)
  - Many-to-One with Job (an application is for one job)

**5. Interview Entity**
- Attributes: _id, company (ObjectId), user (ObjectId), interviewDate, rounds[], overallRating, isAnonymous, createdAt, updatedAt
- Round Sub-document: roundType, questionsAsked[], difficultyRating
- Relationships:
  - Many-to-One with Company (an interview experience is for one company)
  - Many-to-One with User (an interview experience is submitted by one user)

**ENTITY RELATIONSHIP DIAGRAM (ERD) REPRESENTATION:**

```
User (1) ────────< (N) Application
User (1) ────────< (N) Interview
User (N) ────────> (1) Company (for recruiters)

Company (1) ────────< (N) Job
Company (1) ────────< (N) Interview
Company (1) ────────< (N) User (recruiters)

Job (1) ────────< (N) Application
```

**DATA FLOW DIAGRAM (DFD) - LEVEL 0 (CONTEXT DIAGRAM):**

The context diagram shows the system as a single process with external entities.

**External Entities:**
- Student
- Recruiter
- Database (MongoDB)

**Data Flows:**
- Student → System: Registration data, Login credentials, Job application, Interview experience
- System → Student: Job listings, Application status, Interview experiences, Company information
- Recruiter → System: Registration data, Login credentials, Company information, Job postings
- System → Recruiter: Applications, Job listings, Company data
- System ↔ Database: All data operations (CRUD)

**DATA FLOW DIAGRAM (DFD) - LEVEL 1:**

**Processes:**
1. **Authentication Process**
   - Input: User credentials
   - Output: JWT token, User data
   - Data Store: User database

2. **User Management Process**
   - Input: User profile data
   - Output: Updated profile
   - Data Store: User database

3. **Company Management Process**
   - Input: Company information
   - Output: Company data
   - Data Store: Company database

4. **Job Management Process**
   - Input: Job posting data
   - Output: Job listings
   - Data Store: Job database, Company database

5. **Application Management Process**
   - Input: Application data
   - Output: Application status
   - Data Store: Application database, Job database, User database

6. **Interview Experience Process**
   - Input: Interview data
   - Output: Interview experiences
   - Data Store: Interview database, Company database

**USE CASE DIAGRAM:**

**Actors:**
- Student
- Recruiter
- System (automatic processes)

**Student Use Cases:**
- Register Account
- Login
- Browse Jobs
- View Job Details
- Apply for Job
- View Applications
- Submit Interview Experience
- View Interview Experiences
- Update Profile
- Upload Avatar
- Browse Companies
- View Company Details

**Recruiter Use Cases:**
- Register Account
- Login
- Create Company Profile
- Update Company Profile
- Post Job
- Update Job
- Delete Job
- View Applications
- Update Application Status
- View Interview Experiences
- Update Profile
- Upload Avatar

**System Use Cases:**
- Authenticate User
- Hash Password
- Generate JWT Token
- Validate Input
- Handle File Upload
- Manage Sessions

**SEQUENCE DIAGRAM - USER REGISTRATION:**

1. User → Frontend: Fill registration form
2. Frontend → Backend: POST /api/auth/register
3. Backend → Validation: Validate input data
4. Backend → Database: Check if email exists
5. Database → Backend: Return result
6. Backend → bcrypt: Hash password
7. Backend → Database: Create user record
8. Database → Backend: Return user data
9. Backend → JWT: Generate token
10. Backend → Frontend: Return user data + token
11. Frontend → localStorage: Store token
12. Frontend → Router: Redirect to dashboard

**SEQUENCE DIAGRAM - JOB APPLICATION:**

1. Student → Frontend: Click "Apply" button
2. Frontend → Backend: POST /api/applications (with JWT)
3. Backend → Middleware: Verify JWT token
4. Backend → Database: Check if application exists
5. Backend → Database: Create application record
6. Database → Backend: Return application data
7. Backend → Frontend: Return success response
8. Frontend → UI: Update application status

**ACTIVITY DIAGRAM - JOB SEARCH AND APPLICATION FLOW:**

1. Start
2. Student logs in
3. Browse jobs page
4. Apply filters (job type, skills, location)
5. View job listings
6. Select job
7. View job details
8. Decision: Already applied?
   - Yes → Show application status
   - No → Continue
9. Click "Apply"
10. System creates application
11. Application status: "applied"
12. End

**COMPONENT DIAGRAM:**

**Frontend Components:**
- React Application
  - Pages (HomePage, LoginPage, JobsPage, etc.)
  - Components (Navbar, Footer, JobCard, etc.)
  - Context (AuthContext)
  - API Client (Axios)

**Backend Components:**
- Express Server
  - Routes (auth, user, company, job, application, interview)
  - Controllers (business logic)
  - Middleware (authentication, validation)
  - Models (Mongoose schemas)

**Database:**
- MongoDB
  - Collections (users, companies, jobs, applications, interviews)

**DEPLOYMENT DIAGRAM:**

**Development Environment:**
- Developer Machine
  - VS Code
  - Node.js
  - Git

**Production Environment:**
- Frontend Hosting (Vercel/Netlify)
  - React Build
  - Static Files
- Backend Hosting (Railway/Heroku)
  - Node.js Runtime
  - Express Server
- Database Hosting (MongoDB Atlas)
  - MongoDB Cluster
  - Data Storage

**NOTE**: Visual UML diagrams (Class Diagram, Use Case Diagram, Sequence Diagrams, Activity Diagrams, Component Diagrams, Deployment Diagrams) should be created using UML modeling tools such as:
- Draw.io (diagrams.net)
- Lucidchart
- Visual Paradigm
- PlantUML
- Microsoft Visio

These diagrams provide a comprehensive visual representation of the system design, making it easier to understand the system architecture, data flow, and component interactions."""

doc.add_paragraph(uml_dfd_text)

# 4.2 SYSTEM FLOW DIAGRAM
doc.add_heading('4.2 SYSTEM FLOW DIAGRAM', level=2)

system_flow_text = """System flow diagrams illustrate the flow of data and control through the system, showing how different processes interact and how data moves through the system.

**OVERALL SYSTEM FLOW:**

**1. USER REGISTRATION FLOW:**

```
Start
  ↓
User visits Registration Page
  ↓
Select Role (Student/Recruiter)
  ↓
Fill Registration Form
  ↓
Submit Form
  ↓
Frontend Validation
  ↓
[Valid?]
  ├─ No → Show Error Messages → Return to Form
  └─ Yes → Send POST /api/auth/register
         ↓
         Backend Validation
         ↓
         [Email Exists?]
         ├─ Yes → Return Error → Show Error Message
         └─ No → Hash Password
                ↓
                Create User Record
                ↓
                Generate JWT Token
                ↓
                Return User Data + Token
                ↓
                Store Token in localStorage
                ↓
                Redirect to Dashboard
                ↓
End
```

**2. USER LOGIN FLOW:**

```
Start
  ↓
User visits Login Page
  ↓
Enter Email and Password
  ↓
Submit Form
  ↓
Frontend Validation
  ↓
[Valid?]
  ├─ No → Show Error Messages
  └─ Yes → Send POST /api/auth/login
         ↓
         Backend: Find User by Email
         ↓
         [User Exists?]
         ├─ No → Return Error → Show Error Message
         └─ Yes → Compare Password (bcrypt)
                ↓
                [Password Match?]
                ├─ No → Return Error → Show Error Message
                └─ Yes → Generate JWT Token
                       ↓
                       Return User Data + Token
                       ↓
                       Store Token in localStorage
                       ↓
                       Redirect to Dashboard
                       ↓
End
```

**3. JOB BROWSING AND APPLICATION FLOW (STUDENT):**

```
Start
  ↓
Student Logs In
  ↓
Navigate to Jobs Page
  ↓
Load Jobs List (GET /api/jobs)
  ↓
Display Jobs with Filters
  ↓
[Apply Filters?]
  ├─ Yes → Filter Jobs → Update Display
  └─ No → Continue
  ↓
Select Job
  ↓
Navigate to Job Details (GET /api/jobs/:id)
  ↓
Display Job Information
  ↓
[Already Applied?]
  ├─ Yes → Show Application Status
  └─ No → Show "Apply" Button
         ↓
         Click "Apply"
         ↓
         POST /api/applications (with JWT)
         ↓
         Backend: Verify Token
         ↓
         [Authorized?]
         ├─ No → Return 401 Error
         └─ Yes → Check Duplicate Application
                ↓
                [Already Applied?]
                ├─ Yes → Return Error
                └─ No → Create Application Record
                       ↓
                       Set Status: "applied"
                       ↓
                       Return Success
                       ↓
                       Update UI: Show "Applied" Status
                       ↓
End
```

**4. JOB POSTING FLOW (RECRUITER):**

```
Start
  ↓
Recruiter Logs In
  ↓
[Company Exists?]
  ├─ No → Create Company Profile
  │      ↓
  │      Fill Company Form
  │      ↓
  │      POST /api/companies
  │      ↓
  │      Create Company Record
  │      ↓
  │      Link to User
  └─ Yes → Continue
  ↓
Navigate to Post Job Page
  ↓
Fill Job Form (Title, Description, Type, Salary, Skills)
  ↓
Submit Form
  ↓
Frontend Validation
  ↓
[Valid?]
  ├─ No → Show Error Messages
  └─ Yes → POST /api/jobs (with JWT)
         ↓
         Backend: Verify Token & Role
         ↓
         [Authorized?]
         ├─ No → Return 403 Error
         └─ Yes → Create Job Record
                ↓
                Link to Company
                ↓
                Return Job Data
                ↓
                Show Success Message
                ↓
                Redirect to Job List
                ↓
End
```

**5. APPLICATION MANAGEMENT FLOW (RECRUITER):**

```
Start
  ↓
Recruiter Logs In
  ↓
Navigate to Manage Applications
  ↓
Load Applications (GET /api/applications/job/:jobId)
  ↓
Display Applications List
  ↓
Select Application
  ↓
View Application Details
  ↓
[Update Status?]
  ├─ No → End
  └─ Yes → Select New Status
         ↓
         PATCH /api/applications/:id
         ↓
         Backend: Verify Token & Authorization
         ↓
         [Authorized?]
         ├─ No → Return 403 Error
         └─ Yes → Update Application Status
                ↓
                Return Updated Application
                ↓
                Update UI
                ↓
End
```

**6. INTERVIEW EXPERIENCE SUBMISSION FLOW:**

```
Start
  ↓
Student Logs In
  ↓
Navigate to Submit Interview Page
  ↓
Select Company
  ↓
Add Interview Rounds
  ↓
For Each Round:
  - Enter Round Type
  - Enter Questions Asked
  - Set Difficulty Rating
  ↓
Set Overall Rating
  ↓
[Anonymous?]
  ├─ Yes → Set isAnonymous: true
  └─ No → Set isAnonymous: false
  ↓
Submit Form
  ↓
POST /api/interviews (with JWT)
  ↓
Backend: Verify Token
  ↓
[Authorized?]
  ├─ No → Return 401 Error
  └─ Yes → Create Interview Record
         ↓
         Link to Company and User
         ↓
         Return Success
         ↓
         Show Success Message
         ↓
         Redirect to Dashboard
         ↓
End
```

**7. PROFILE UPDATE FLOW:**

```
Start
  ↓
User Logs In
  ↓
Navigate to Profile Page
  ↓
Load Current Profile (GET /api/users/profile)
  ↓
Display Profile Form
  ↓
[Update Information?]
  ├─ No → End
  └─ Yes → Modify Form Fields
         ↓
         [Upload Avatar?]
         ├─ Yes → Select Image File
         │      ↓
         │      POST /api/users/avatar (multipart/form-data)
         │      ↓
         │      Backend: Save File
         │      ↓
         │      Update User Avatar URL
         └─ No → Continue
         ↓
         Submit Profile Update
         ↓
         PUT /api/users/profile (with JWT)
         ↓
         Backend: Verify Token
         ↓
         Update User Record
         ↓
         Return Updated User Data
         ↓
         Update UI
         ↓
End
```

**SYSTEM ARCHITECTURE FLOW:**

```
Client (Browser)
  ↓
React Frontend Application
  ↓
React Router (Navigation)
  ↓
API Calls (Axios)
  ↓
Express Backend Server
  ↓
Middleware (Auth, Validation)
  ↓
Controllers (Business Logic)
  ↓
Mongoose Models
  ↓
MongoDB Database
```

**AUTHENTICATION FLOW:**

```
Request with JWT Token
  ↓
Auth Middleware
  ↓
Extract Token from Header
  ↓
[Token Present?]
  ├─ No → Return 401 Unauthorized
  └─ Yes → Verify Token (JWT.verify)
         ↓
         [Token Valid?]
         ├─ No → Return 401 Unauthorized
         └─ Yes → Extract User Info
                ↓
                Attach to Request Object
                ↓
                Continue to Route Handler
```

**ERROR HANDLING FLOW:**

```
Error Occurs
  ↓
Error Handler Middleware
  ↓
Identify Error Type
  ↓
[Multer Error?]
  ├─ Yes → Return 400: File Upload Error
  └─ No → [JWT Error?]
          ├─ Yes → Return 401: Authentication Error
          └─ No → [Validation Error?]
                 ├─ Yes → Return 400: Validation Error
                 └─ No → Return 500: Internal Server Error
                        ↓
                        Log Error (Development)
                        ↓
                        Return Error Response
                        ↓
                        Frontend: Display Error Message
```

These system flow diagrams provide a clear understanding of how data and control flow through the Ahmedabad Career Hub system, illustrating the sequence of operations for various user interactions and system processes."""

doc.add_paragraph(system_flow_text)

# 4.3 DATA DICTIONARY
doc.add_heading('4.3 DATA DICTIONARY', level=2)

data_dictionary_text = """The data dictionary provides detailed information about all data elements, their attributes, types, constraints, and relationships in the Ahmedabad Career Hub system.

**USER ENTITY:**

**Table/Collection Name:** users

**Attributes:**
1. **_id**
   - Type: ObjectId (MongoDB)
   - Description: Primary key, unique identifier
   - Constraints: Auto-generated, Required, Unique

2. **name**
   - Type: String
   - Description: User's full name
   - Constraints: Required, Minimum 2 characters
   - Example: "John Doe"

3. **email**
   - Type: String
   - Description: User's email address
   - Constraints: Required, Unique, Valid email format
   - Example: "john.doe@example.com"

4. **password**
   - Type: String (hashed)
   - Description: User's password (hashed using bcrypt)
   - Constraints: Required, Minimum 8 characters, Must contain uppercase, lowercase, and number
   - Example: "$2b$10$hashedPasswordString"

5. **role**
   - Type: String (Enum)
   - Description: User role (student or recruiter)
   - Constraints: Required, Enum: ["student", "recruiter"]
   - Example: "student"

6. **college**
   - Type: String
   - Description: College name (for students)
   - Constraints: Optional, Minimum 2 characters (if provided)
   - Example: "Gujarat University"

7. **companyName**
   - Type: String
   - Description: Company name entered during registration (for recruiters)
   - Constraints: Optional
   - Example: "Tech Solutions Inc."

8. **company**
   - Type: ObjectId (Reference)
   - Description: Reference to Company document (for recruiters)
   - Constraints: Optional, References Company._id
   - Example: ObjectId("507f1f77bcf86cd799439011")

9. **avatarUrl**
   - Type: String
   - Description: URL/path to user's avatar image
   - Constraints: Optional
   - Example: "/uploads/avatars/user123.jpg"

10. **createdAt**
    - Type: Date
    - Description: Timestamp when user was created
    - Constraints: Auto-generated
    - Example: 2024-01-15T10:30:00Z

11. **updatedAt**
    - Type: Date
    - Description: Timestamp when user was last updated
    - Constraints: Auto-updated
    - Example: 2024-01-20T14:45:00Z

**COMPANY ENTITY:**

**Table/Collection Name:** companies

**Attributes:**
1. **_id**
   - Type: ObjectId
   - Description: Primary key
   - Constraints: Auto-generated, Required, Unique

2. **companyName**
   - Type: String
   - Description: Company name
   - Constraints: Required, Unique
   - Example: "Tech Solutions Inc."

3. **description**
   - Type: String
   - Description: Company description
   - Constraints: Required
   - Example: "A leading software development company..."

4. **website**
   - Type: String
   - Description: Company website URL
   - Constraints: Optional, Valid URL format
   - Example: "https://www.techsolutions.com"

5. **logoUrl**
   - Type: String
   - Description: URL/path to company logo image
   - Constraints: Optional
   - Example: "/uploads/logos/company123.jpg"

6. **address**
   - Type: String
   - Description: Company address
   - Constraints: Required
   - Example: "123 Tech Street, Ahmedabad, Gujarat"

7. **techStack**
   - Type: Array of Strings
   - Description: List of technologies used by company
   - Constraints: Optional, Array of strings
   - Example: ["React", "Node.js", "MongoDB", "Python"]

8. **createdAt**
   - Type: Date
   - Description: Timestamp when company was created
   - Constraints: Auto-generated

9. **updatedAt**
   - Type: Date
   - Description: Timestamp when company was last updated
   - Constraints: Auto-updated

**JOB ENTITY:**

**Table/Collection Name:** jobs

**Attributes:**
1. **_id**
   - Type: ObjectId
   - Description: Primary key
   - Constraints: Auto-generated, Required, Unique

2. **title**
   - Type: String
   - Description: Job title
   - Constraints: Required
   - Example: "Full Stack Developer"

3. **description**
   - Type: String
   - Description: Detailed job description
   - Constraints: Required
   - Example: "We are looking for a skilled full stack developer..."

4. **jobType**
   - Type: String (Enum)
   - Description: Type of job (internship or full-time)
   - Constraints: Required, Enum: ["internship", "full-time"]
   - Example: "full-time"

5. **salaryStipend**
   - Type: String
   - Description: Salary or stipend amount
   - Constraints: Required
   - Example: "₹50,000 - ₹80,000" or "₹15,000/month"

6. **location**
   - Type: String
   - Description: Job location
   - Constraints: Optional, Default: "Ahmedabad"
   - Example: "Ahmedabad"

7. **requiredSkills**
   - Type: Array of Strings
   - Description: List of required skills
   - Constraints: Optional, Array of strings
   - Example: ["React", "Node.js", "MongoDB", "JavaScript"]

8. **company**
   - Type: ObjectId (Reference)
   - Description: Reference to Company document
   - Constraints: Required, References Company._id
   - Example: ObjectId("507f1f77bcf86cd799439011")

9. **createdAt**
   - Type: Date
   - Description: Timestamp when job was posted
   - Constraints: Auto-generated

10. **updatedAt**
    - Type: Date
    - Description: Timestamp when job was last updated
    - Constraints: Auto-updated

**APPLICATION ENTITY:**

**Table/Collection Name:** applications

**Attributes:**
1. **_id**
   - Type: ObjectId
   - Description: Primary key
   - Constraints: Auto-generated, Required, Unique

2. **user**
   - Type: ObjectId (Reference)
   - Description: Reference to User document (applicant)
   - Constraints: Required, References User._id
   - Example: ObjectId("507f1f77bcf86cd799439012")

3. **job**
   - Type: ObjectId (Reference)
   - Description: Reference to Job document
   - Constraints: Required, References Job._id
   - Example: ObjectId("507f1f77bcf86cd799439013")

4. **appliedAt**
   - Type: Date
   - Description: Timestamp when application was submitted
   - Constraints: Auto-generated, Default: Date.now
   - Example: 2024-01-20T10:00:00Z

5. **status**
   - Type: String (Enum)
   - Description: Application status
   - Constraints: Required, Default: "applied", Enum: ["applied", "reviewed", "shortlisted", "rejected", "hired"]
   - Example: "applied"

**INTERVIEW ENTITY:**

**Table/Collection Name:** interviews

**Attributes:**
1. **_id**
   - Type: ObjectId
   - Description: Primary key
   - Constraints: Auto-generated, Required, Unique

2. **company**
   - Type: ObjectId (Reference)
   - Description: Reference to Company document
   - Constraints: Required, References Company._id
   - Example: ObjectId("507f1f77bcf86cd799439011")

3. **user**
   - Type: ObjectId (Reference)
   - Description: Reference to User document (interview submitter)
   - Constraints: Required, References User._id
   - Example: ObjectId("507f1f77bcf86cd799439012")

4. **interviewDate**
   - Type: Date
   - Description: Date of interview
   - Constraints: Optional, Default: Date.now
   - Example: 2024-01-15T09:00:00Z

5. **rounds**
   - Type: Array of Objects
   - Description: Array of interview rounds
   - Constraints: Required, Array of round objects
   - Structure:
     - roundType: String (Required) - e.g., "Technical", "HR", "Coding Test"
     - questionsAsked: Array of Strings (Required) - List of questions
     - difficultyRating: Number (Required, 1-5) - Difficulty rating
   - Example: [{
       roundType: "Technical",
       questionsAsked: ["What is React?", "Explain closures"],
       difficultyRating: 4
     }]

6. **overallRating**
   - Type: Number
   - Description: Overall interview experience rating
   - Constraints: Required, Min: 1, Max: 5
   - Example: 4

7. **isAnonymous**
   - Type: Boolean
   - Description: Whether the interview experience is posted anonymously
   - Constraints: Optional, Default: true
   - Example: true

8. **createdAt**
   - Type: Date
   - Description: Timestamp when interview experience was created
   - Constraints: Auto-generated

9. **updatedAt**
   - Type: Date
   - Description: Timestamp when interview experience was last updated
   - Constraints: Auto-updated

**INDEXES:**

**User Collection:**
- email: Unique index

**Company Collection:**
- companyName: Unique index

**Application Collection:**
- user: Index
- job: Index
- Compound index: (user, job) - Unique (prevents duplicate applications)

**Interview Collection:**
- company: Index
- user: Index

**RELATIONSHIPS SUMMARY:**

1. **User → Application**: One-to-Many (1:N)
2. **User → Interview**: One-to-Many (1:N)
3. **User → Company**: Many-to-One (N:1) - for recruiters
4. **Company → User**: One-to-Many (1:N) - recruiters
5. **Company → Job**: One-to-Many (1:N)
6. **Company → Interview**: One-to-Many (1:N)
7. **Job → Application**: One-to-Many (1:N)

**DATA TYPES SUMMARY:**

- **ObjectId**: MongoDB unique identifier (24 hex characters)
- **String**: Text data
- **Number**: Numeric data
- **Boolean**: True/false values
- **Date**: Timestamp data
- **Array**: List of values
- **Object**: Nested document structure
- **Enum**: Restricted set of string values

This data dictionary provides a comprehensive reference for all data elements in the Ahmedabad Career Hub system, ensuring consistency and clarity in system design and implementation."""

doc.add_paragraph(data_dictionary_text)

# 4.4 USER INTERFACE
doc.add_heading('4.4 USER INTERFACE', level=2)

ui_text = """The user interface design focuses on creating an intuitive, responsive, and user-friendly experience for both students and recruiters. The interface is built using React components with TailwindCSS for styling.

**UI DESIGN PRINCIPLES:**

1. **Simplicity**: Clean and uncluttered design
2. **Consistency**: Uniform design patterns across all pages
3. **Responsiveness**: Works seamlessly on desktop, tablet, and mobile devices
4. **Accessibility**: Clear navigation and readable content
5. **User-Centric**: Designed based on user needs and workflows

**PAGE STRUCTURE:**

**1. HOME PAGE (/):**
- **Purpose**: Landing page and introduction
- **Components**:
  - Hero section with animated background
  - Main heading: "Ahmedabad Career Hub"
  - Tagline and description
  - Call-to-action buttons (Register, Browse Jobs)
  - Statistics section (Active Jobs, Companies, Students, Success Rate)
  - Features section (6 feature cards)
  - Testimonial section
  - Final CTA section
- **Design Elements**:
  - Gradient backgrounds
  - Animated floating elements
  - Hover effects
  - Responsive grid layout

**2. LOGIN PAGE (/login):**
- **Purpose**: User authentication
- **Components**:
  - Login form (Email, Password)
  - Submit button
  - Link to registration page
  - Error message display
- **Form Fields**:
  - Email: Text input, required, email validation
  - Password: Password input, required
- **Validation**: Client-side and server-side validation

**3. REGISTRATION PAGE (/register):**
- **Purpose**: New user registration
- **Components**:
  - Role selection (Student/Recruiter)
  - Registration form (dynamic based on role)
  - Submit button
  - Link to login page
  - Error message display
- **Form Fields (Student)**:
  - Name: Text input, required
  - Email: Text input, required, email validation
  - Password: Password input, required, strength validation
  - College: Text input, required
- **Form Fields (Recruiter)**:
  - Name: Text input, required
  - Email: Text input, required, email validation
  - Password: Password input, required, strength validation
  - Company Name: Text input, optional

**4. DASHBOARD PAGE (/dashboard):**
- **Purpose**: Student dashboard with overview
- **Components**:
  - Welcome message with user name
  - Recent applications section (shows 3 most recent)
  - Quick action buttons:
    - View All Applications
    - Submit Interview Experience
    - Browse Jobs
    - Browse Companies
  - Application cards with status indicators
- **Design**: Card-based layout, color-coded status badges

**5. RECRUITER DASHBOARD (/recruiter-dashboard):**
- **Purpose**: Recruiter dashboard with job management
- **Components**:
  - Welcome message
  - Posted jobs list
  - Quick actions:
    - Post New Job
    - Manage Applications
    - View My Company
  - Job statistics
- **Design**: Table/card layout for jobs

**6. JOBS PAGE (/jobs):**
- **Purpose**: Browse and search jobs
- **Components**:
  - Search bar
  - Filter options:
    - Job Type (Internship/Full-time)
    - Skills filter
    - Location filter
  - Job listings (paginated)
  - Job cards showing:
    - Job title
    - Company name
    - Job type
    - Location
    - Required skills
    - Salary/Stipend
  - Pagination controls
- **Design**: Grid layout, responsive cards

**7. JOB DETAIL PAGE (/jobs/:id):**
- **Purpose**: Detailed job information
- **Components**:
  - Job title and company name
  - Full job description
  - Job details (type, location, salary)
  - Required skills list
  - Company information card
  - Apply button (for students)
  - Application status (if already applied)
- **Design**: Single column layout, prominent CTA button

**8. COMPANIES PAGE (/companies):**
- **Purpose**: Browse companies
- **Components**:
  - Company search bar
  - Company listings (paginated)
  - Company cards showing:
    - Company name
    - Logo
    - Description preview
    - Tech stack tags
    - Location
  - Pagination controls
- **Design**: Grid layout, company cards

**9. COMPANY DETAIL PAGE (/companies/:id):**
- **Purpose**: Detailed company information
- **Components**:
  - Company name and logo
  - Full company description
  - Company details (website, address)
  - Tech stack display
  - Available jobs list
  - Interview experiences section
- **Design**: Two-column layout, information-rich

**10. APPLICATIONS PAGE (/applications):**
- **Purpose**: Student's application history
- **Components**:
  - Application list (paginated)
  - Application cards showing:
    - Job title
    - Company name
    - Application status
    - Applied date
    - Status badge (color-coded)
  - Filter by status option
  - Pagination controls
- **Design**: List layout, status indicators

**11. MANAGE APPLICATIONS PAGE (/manage-applications):**
- **Purpose**: Recruiter's application management
- **Components**:
  - Job selector
  - Applications list for selected job
  - Application cards with:
    - Applicant name
    - Application date
    - Current status
    - Status update dropdown
  - Status update functionality
- **Design**: Table/card layout, action buttons

**12. POST JOB PAGE (/post-job):**
- **Purpose**: Create new job posting
- **Components**:
  - Job posting form:
    - Job title (required)
    - Job description (required, textarea)
    - Job type (Internship/Full-time, required)
    - Salary/Stipend (required)
    - Location (required)
    - Required skills (multi-select or tags)
  - Submit button
  - Cancel button
- **Design**: Form layout, validation feedback

**13. SUBMIT INTERVIEW PAGE (/submit-interview):**
- **Purpose**: Submit interview experience
- **Components**:
  - Company selector
  - Interview date picker
  - Rounds section (dynamic):
    - Round type
    - Questions asked (textarea)
    - Difficulty rating (1-5)
    - Add/Remove round buttons
  - Overall rating (1-5)
  - Anonymous option (checkbox)
  - Submit button
- **Design**: Multi-section form, dynamic fields

**14. PROFILE PAGE (/profile):**
- **Purpose**: User profile management
- **Components**:
  - Profile information form:
    - Name (editable)
    - Email (read-only)
    - College (for students, editable)
    - Company (for recruiters, read-only)
  - Avatar upload section:
    - Current avatar display
    - Upload button
    - Image preview
  - Update button
- **Design**: Form layout, image upload area

**COMMON COMPONENTS:**

**1. NAVBAR:**
- Logo/Brand name
- Navigation links (Home, Jobs, Companies)
- User menu (when logged in):
  - Dashboard
  - Profile
  - Logout
- Login/Register buttons (when not logged in)
- Responsive hamburger menu (mobile)

**2. FOOTER:**
- Copyright information
- Links to important pages
- Social media links (optional)
- Contact information

**3. JOB CARD:**
- Reusable component for job display
- Shows: Title, Company, Type, Location, Skills, Salary
- Clickable to navigate to job details

**4. COMPANY CARD:**
- Reusable component for company display
- Shows: Name, Logo, Description, Tech Stack
- Clickable to navigate to company details

**5. BUTTON COMPONENT:**
- Primary button (blue gradient)
- Secondary button (outline)
- Loading state
- Disabled state

**6. INPUT COMPONENT:**
- Text input
- Password input
- Textarea
- Select dropdown
- Error message display

**7. PROTECTED ROUTE:**
- Wrapper component for authenticated routes
- Redirects to login if not authenticated
- Shows loading state during authentication check

**COLOR SCHEME:**

- **Primary**: Blue (#3B82F6, #2563EB)
- **Secondary**: Purple (#9333EA)
- **Accent**: Cyan (#06B6D4)
- **Success**: Green (#10B981)
- **Error**: Red (#EF4444)
- **Warning**: Yellow (#F59E0B)
- **Background**: White, Light Gray (#F9FAFB)
- **Text**: Dark Gray (#1F2937), Gray (#6B7280)

**TYPOGRAPHY:**

- **Headings**: Bold, Large sizes (2xl, 3xl, 4xl, 5xl)
- **Body Text**: Regular, Medium sizes (base, lg)
- **Font Family**: System fonts (sans-serif)

**RESPONSIVE DESIGN:**

- **Desktop**: Full layout, multi-column grids
- **Tablet**: Adjusted grid, stacked elements
- **Mobile**: Single column, hamburger menu, touch-friendly buttons

**UI FEATURES:**

- **Loading States**: Spinners and skeleton loaders
- **Error Messages**: Toast notifications (react-hot-toast)
- **Success Messages**: Toast notifications
- **Form Validation**: Real-time validation feedback
- **Hover Effects**: Interactive elements with hover states
- **Animations**: Smooth transitions and animations
- **Pagination**: Page numbers and navigation
- **Search**: Real-time search with debouncing
- **Filters**: Collapsible filter panels

**ACCESSIBILITY:**

- Semantic HTML elements
- ARIA labels where needed
- Keyboard navigation support
- Focus indicators
- Alt text for images
- Color contrast compliance

This user interface design ensures a modern, intuitive, and accessible experience for all users of the Ahmedabad Career Hub platform."""

doc.add_paragraph(ui_text)

# 4.5 SYSTEM NAVIGATION
doc.add_heading('4.5 SYSTEM NAVIGATION', level=2)

navigation_text = """System navigation defines how users move through the application, access different features, and interact with the system. The navigation structure is designed to be intuitive and role-specific.

**NAVIGATION ARCHITECTURE:**

**PUBLIC NAVIGATION (Unauthenticated Users):**

```
Home Page (/)
  ├─ Login (/login)
  ├─ Register (/register)
  ├─ Jobs (/jobs)
  │   └─ Job Details (/jobs/:id)
  └─ Companies (/companies)
      └─ Company Details (/companies/:id)
```

**STUDENT NAVIGATION (Authenticated Students):**

```
Dashboard (/dashboard)
  ├─ Applications (/applications)
  ├─ Submit Interview (/submit-interview)
  ├─ Jobs (/jobs)
  │   └─ Job Details (/jobs/:id)
  │       └─ Apply to Job
  ├─ Companies (/companies)
  │   └─ Company Details (/companies/:id)
  └─ Profile (/profile)
```

**RECRUITER NAVIGATION (Authenticated Recruiters):**

```
Recruiter Dashboard (/recruiter-dashboard)
  ├─ Post Job (/post-job)
  ├─ Edit Job (/edit-job/:id)
  ├─ Manage Applications (/manage-applications)
  ├─ Create Company (/create-company)
  ├─ My Company (/my-company)
  ├─ Jobs (/jobs)
  │   └─ Job Details (/jobs/:id)
  ├─ Companies (/companies)
  │   └─ Company Details (/companies/:id)
  └─ Profile (/profile)
```

**NAVIGATION COMPONENTS:**

**1. NAVBAR NAVIGATION:**

**Public Navbar:**
- Logo (links to Home)
- Home
- Jobs
- Companies
- Login (button)
- Register (button)

**Authenticated Navbar:**
- Logo (links to Home)
- Home
- Jobs
- Companies
- User Menu (dropdown):
  - Dashboard (role-specific)
  - Profile
  - Logout

**2. DASHBOARD NAVIGATION:**

**Student Dashboard Quick Actions:**
- View All Applications → /applications
- Submit Interview Experience → /submit-interview
- Browse Jobs → /jobs
- Browse Companies → /companies

**Recruiter Dashboard Quick Actions:**
- Post New Job → /post-job
- Manage Applications → /manage-applications
- View My Company → /my-company
- Create Company → /create-company (if no company)

**3. BREADCRUMB NAVIGATION:**

**Example Paths:**
- Home > Jobs > Job Details
- Home > Companies > Company Details
- Dashboard > Applications
- Dashboard > Profile

**4. FOOTER NAVIGATION:**

- Home
- About (if applicable)
- Contact (if applicable)
- Privacy Policy (if applicable)
- Terms of Service (if applicable)

**ROUTING STRUCTURE:**

**Public Routes:**
- `/` - HomePage
- `/login` - LoginPage
- `/register` - RegisterPage
- `/jobs` - JobsPage
- `/jobs/:id` - JobDetailPage
- `/companies` - CompaniesPage
- `/companies/:id` - CompanyDetailPage

**Protected Routes (Require Authentication):**
- `/dashboard` - DashboardPage (Students)
- `/recruiter-dashboard` - RecruiterDashboardPage (Recruiters)
- `/applications` - ApplicationsPage (Students)
- `/submit-interview` - SubmitInterviewPage (Students)
- `/post-job` - PostJobPage (Recruiters)
- `/edit-job/:id` - EditJobPage (Recruiters)
- `/manage-applications` - ManageApplicationsPage (Recruiters)
- `/create-company` - CreateCompanyPage (Recruiters)
- `/my-company` - MyCompanyPage (Recruiters)
- `/profile` - ProfilePage (All authenticated users)

**NAVIGATION FLOWS:**

**1. REGISTRATION FLOW:**
```
Home → Register → (Select Role) → Fill Form → Submit → Auto-login → Dashboard
```

**2. LOGIN FLOW:**
```
Home → Login → Enter Credentials → Submit → Verify → Dashboard
```

**3. JOB APPLICATION FLOW (STUDENT):**
```
Dashboard → Jobs → Select Job → Job Details → Apply → Application Created → Applications Page
```

**4. JOB POSTING FLOW (RECRUITER):**
```
Recruiter Dashboard → Post Job → Fill Form → Submit → Job Created → Job List
```

**5. INTERVIEW SUBMISSION FLOW (STUDENT):**
```
Dashboard → Submit Interview → Select Company → Fill Form → Submit → Interview Submitted → Dashboard
```

**6. APPLICATION MANAGEMENT FLOW (RECRUITER):**
```
Recruiter Dashboard → Manage Applications → Select Job → View Applications → Update Status → Status Updated
```

**NAVIGATION PATTERNS:**

**1. HIERARCHICAL NAVIGATION:**
- Top-level pages accessible from navbar
- Sub-pages accessible from parent pages
- Breadcrumbs show navigation path

**2. CONTEXTUAL NAVIGATION:**
- Action buttons on detail pages
- Quick actions on dashboard
- Related links on detail pages

**3. MODAL NAVIGATION:**
- Confirmation dialogs
- Form modals (if applicable)
- Alert messages

**4. TAB NAVIGATION:**
- Profile tabs (if applicable)
- Dashboard sections (if applicable)

**NAVIGATION STATES:**

**1. ACTIVE STATE:**
- Current page highlighted in navbar
- Active link styling
- Breadcrumb indication

**2. DISABLED STATE:**
- Unavailable actions disabled
- Visual indication of disabled state
- Tooltip explaining why disabled

**3. LOADING STATE:**
- Navigation during data loading
- Loading indicators
- Disabled navigation during operations

**4. ERROR STATE:**
- Error pages (404, 500)
- Error messages
- Retry options

**ROLE-BASED NAVIGATION:**

**Student-Specific Routes:**
- `/dashboard` - Student dashboard
- `/applications` - Application history
- `/submit-interview` - Interview submission

**Recruiter-Specific Routes:**
- `/recruiter-dashboard` - Recruiter dashboard
- `/post-job` - Job posting
- `/manage-applications` - Application management
- `/create-company` - Company creation
- `/my-company` - Company management
- `/edit-job/:id` - Job editing

**NAVIGATION GUARDS:**

**1. AUTHENTICATION GUARD:**
- Protected routes require authentication
- Redirects to login if not authenticated
- Preserves intended destination after login

**2. ROLE GUARD:**
- Role-specific routes check user role
- Redirects to appropriate dashboard
- Prevents unauthorized access

**3. DATA GUARD:**
- Checks if required data exists
- Redirects if data not found (404)
- Handles missing resources gracefully

**MOBILE NAVIGATION:**

**Hamburger Menu:**
- Collapsible menu on mobile
- Icon-based navigation
- Slide-out menu (if applicable)

**Bottom Navigation (if applicable):**
- Fixed bottom bar on mobile
- Quick access to main sections
- Icon-based navigation

**NAVIGATION ENHANCEMENTS:**

**1. SEARCH NAVIGATION:**
- Global search bar
- Search results page
- Quick search suggestions

**2. FILTER NAVIGATION:**
- Filter panels
- Applied filters display
- Clear filters option

**3. PAGINATION NAVIGATION:**
- Page numbers
- Previous/Next buttons
- Jump to page (if applicable)

**4. BREADCRUMB NAVIGATION:**
- Shows current location
- Clickable path segments
- Home link always available

**NAVIGATION ACCESSIBILITY:**

- Keyboard navigation support
- Focus indicators
- ARIA labels
- Screen reader compatibility
- Skip navigation links

**NAVIGATION PERFORMANCE:**

- Lazy loading of routes
- Code splitting
- Prefetching (if applicable)
- Optimized route transitions

This navigation structure ensures users can easily access all features of the Ahmedabad Career Hub system, with clear paths for different user roles and intuitive navigation patterns throughout the application."""

doc.add_paragraph(navigation_text)

# Save the updated document
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("Chapter 4 has been added successfully to the document!")

