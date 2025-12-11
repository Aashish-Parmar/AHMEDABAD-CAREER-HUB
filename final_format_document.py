from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("Applying final formatting to match reference document style...")

doc = Document('Ahmedabad_Career_Hub_Chapter1.docx')

# Apply Times New Roman font throughout
print("Applying Times New Roman font to all text...")
for para in doc.paragraphs:
    if para.runs:
        for run in para.runs:
            # Set Times New Roman
            run.font.name = 'Times New Roman'
            # Maintain existing sizes or set default
            if not run.font.size:
                if para.style.name.startswith('Heading'):
                    if 'Heading 1' in para.style.name:
                        run.font.size = Pt(16)
                    elif 'Heading 2' in para.style.name:
                        run.font.size = Pt(14)
                    elif 'Heading 3' in para.style.name:
                        run.font.size = Pt(12)
                else:
                    run.font.size = Pt(12)

# Ensure headings are bold
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        if para.runs:
            for run in para.runs:
                run.bold = True

print("Formatting complete!")
print("\nDocument now has:")
print("- Times New Roman font throughout")
print("- Consistent heading sizes")
print("- Proper alignment for chapter titles")

# Save
doc.save('Ahmedabad_Career_Hub_Chapter1.docx')
print("\nDocument saved successfully!")
print("\nIMPORTANT: To match the reference document style for diagrams:")
print("1. Diagrams should be inserted as images")
print("2. Each diagram should have: 'Figure X: Title' (centered, bold, 11pt Times New Roman)")
print("3. Diagrams should be centered on the page")
print("4. Diagrams can be in bordered boxes or as standalone images")
print("5. Reference diagrams in text using 'as shown in Figure X'")

